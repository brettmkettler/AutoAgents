# filename: create_readme_doc.py
from docx import Document

def create_readme_doc():
    doc = Document()
    doc.add_heading('Readme Document for Internet Search Script', 0)

    doc.add_paragraph("This script searches the internet for a given query. It uses the 'googlesearch-python' library to search for a query on Google and prints out the URLs of the search results.")
    doc.add_paragraph("To use the script, simply run it with Python and it will print out the URLs of the search results for the query 'Python scripting'.")
    doc.add_paragraph("Feel free to modify and use it for your needs.")

    doc.save("readme.docx")

create_readme_doc()